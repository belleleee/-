"""Human-readable report renderers for delivery-friendly artifacts."""

from __future__ import annotations

from html import escape
from pathlib import Path
import subprocess
from risk_himate.app.core.schemas import RiskReport, RiskReportDetail


LEVEL_LABELS = {
    "low": "低",
    "medium": "中",
    "high": "高",
}

TREND_LABELS = {
    None: "暂无",
    "up": "上升",
    "down": "下降",
    "stable": "稳定",
}


def render_risk_report_markdown(report: RiskReport) -> str:
    top_risks = [_ensure_detail(detail) for detail in report.top3_risks]
    risk_details = [_ensure_detail(detail) for detail in report.risk_details]
    human_review_items = [_ensure_detail(detail) for detail in report.human_review_items]
    lines: list[str] = [
        f"# {report.company} 风险分析报告",
        "",
        "## 一、执行摘要",
        "",
        f"- 评估时间：`{report.timestamp}`",
        f"- 综合风险等级：`{LEVEL_LABELS.get(report.overall_risk_level, report.overall_risk_level)}`",
        f"- 综合风险分数：`{report.overall_score}`",
        f"- 模型综合置信度：`{report.confidence}`",
        f"- 风险趋势：`{TREND_LABELS.get(report.trend, report.trend)}`",
        f"- 重点风险类别：`{', '.join(report.top_categories) if report.top_categories else '暂无'}`",
        "",
        _build_overall_narrative(report),
        "",
        "## 二、置信度与红线校验",
        "",
        *_build_confidence_lines(report),
        "",
        "## 三、重点风险概览",
        "",
    ]

    if top_risks:
        for index, detail in enumerate(top_risks, start=1):
            lines.extend(_render_top_risk_markdown(index, detail))
    else:
        lines.append("当前未识别出明确风险。")

    lines.extend(
        [
            "",
            "## 四、按生命周期阶段的风险分析与解决方案",
            "",
        ]
    )

    if report.lifecycle_stage_groups:
        running_index = 1
        for group in report.lifecycle_stage_groups:
            lines.append(f"### {group.stage}")
            lines.append("")
            lines.append(group.summary)
            lines.append("")
            for hint in group.propagation_hints:
                lines.append(f"- {hint}")
            if group.propagation_hints:
                lines.append("")
            for detail in group.risk_details:
                lines.extend(_render_detail_markdown(running_index, detail))
                running_index += 1
    elif risk_details:
        for index, detail in enumerate(risk_details, start=1):
            lines.extend(_render_detail_markdown(index, detail))
    else:
        lines.append("当前无风险明细。")

    lines.extend(
        [
            "",
            "## 五、建议行动计划",
            "",
        ]
    )
    lines.extend(_build_action_plan_lines(report))

    lines.extend(
        [
            "",
            "## 六、人工复核建议",
            "",
        ]
    )
    if human_review_items:
        for item in human_review_items:
            lines.append(
                f"- `{item.category} / {item.subtype}`：建议人工复核，当前置信度为 `{item.confidence}`。"
            )
    else:
        lines.append("当前无必须人工复核的条目。")

    return "\n".join(lines).strip() + "\n"


def render_risk_report_html(report: RiskReport) -> str:
    risk_details = [_ensure_detail(detail) for detail in report.risk_details]
    body = [
        "<!DOCTYPE html>",
        '<html lang="zh-CN">',
        "<head>",
        '<meta charset="utf-8">',
        "<title>Risk-HiMATE Report</title>",
        "<style>",
        "body { font-family: -apple-system, BlinkMacSystemFont, 'PingFang SC', 'Microsoft YaHei', sans-serif; margin: 40px; color: #1f2937; line-height: 1.7; }",
        "h1, h2, h3 { color: #111827; }",
        ".meta { background: #f8fafc; border: 1px solid #e5e7eb; padding: 16px; border-radius: 10px; }",
        ".risk-card { border: 1px solid #e5e7eb; border-radius: 10px; padding: 16px; margin: 16px 0; }",
        ".high { border-left: 6px solid #dc2626; }",
        ".medium { border-left: 6px solid #d97706; }",
        ".low { border-left: 6px solid #2563eb; }",
        "code { background: #f3f4f6; padding: 2px 6px; border-radius: 4px; }",
        "ul { padding-left: 20px; }",
        "</style>",
        "</head>",
        "<body>",
        f"<h1>{escape(report.company)} 风险分析报告</h1>",
        '<div class="meta">',
        f"<p><strong>评估时间：</strong>{escape(report.timestamp)}</p>",
        f"<p><strong>综合风险等级：</strong>{escape(LEVEL_LABELS.get(report.overall_risk_level, report.overall_risk_level))}</p>",
        f"<p><strong>综合风险分数：</strong>{report.overall_score}</p>",
        f"<p><strong>模型综合置信度：</strong>{report.confidence}</p>",
        f"<p><strong>风险趋势：</strong>{escape(TREND_LABELS.get(report.trend, report.trend or '暂无'))}</p>",
        f"<p><strong>重点风险类别：</strong>{escape(', '.join(report.top_categories) if report.top_categories else '暂无')}</p>",
        "</div>",
        "<h2>一、执行摘要</h2>",
        f"<p>{escape(_build_overall_narrative(report))}</p>",
        "<h2>二、置信度与红线校验</h2>",
    ]
    for line in _build_confidence_lines(report):
        if line.startswith("- "):
            body.append(f"<p>{escape(line[2:])}</p>")
        else:
            body.append(f"<p>{escape(line)}</p>")
    body.append("<h2>三、按生命周期阶段的风险分析与解决方案</h2>")

    if report.lifecycle_stage_groups:
        running_index = 1
        for group in report.lifecycle_stage_groups:
            body.append(f"<h3>{escape(group.stage)}</h3>")
            body.append(f"<p>{escape(group.summary)}</p>")
            for hint in group.propagation_hints:
                body.append(f"<p><strong>传导提示：</strong>{escape(hint)}</p>")
            for detail in group.risk_details:
                body.extend(_render_detail_html(running_index, detail))
                running_index += 1
    elif risk_details:
        for index, detail in enumerate(risk_details, start=1):
            body.extend(_render_detail_html(index, detail))
    else:
        body.append("<p>当前无风险明细。</p>")

    body.extend(
        [
            "<h2>四、建议行动计划</h2>",
            "<ul>",
        ]
    )
    for line in _build_action_plan_lines(report):
        if line.startswith("- "):
            body.append(f"<li>{escape(line[2:])}</li>")
    body.extend(["</ul>", "</body>", "</html>"])
    return "\n".join(body)


def render_risk_report_rtf(report: RiskReport) -> str:
    markdown = render_risk_report_markdown(report)
    escaped = (
        markdown.replace("\\", "\\\\")
        .replace("{", "\\{")
        .replace("}", "\\}")
    )
    escaped = escaped.replace("\n", "\\par\n")
    return "{\\rtf1\\ansi\\deff0\n" + escaped + "\n}"


def write_rendered_report(report: RiskReport, output_path: Path) -> None:
    suffix = output_path.suffix.lower()
    if suffix == ".md":
        output_path.write_text(render_risk_report_markdown(report), encoding="utf-8")
        return
    if suffix == ".rtf":
        output_path.write_text(render_risk_report_rtf(report), encoding="utf-8")
        return
    if suffix in {".html", ".htm"}:
        output_path.write_text(render_risk_report_html(report), encoding="utf-8")
        return
    if suffix == ".docx":
        _write_docx_report(report, output_path)
        return
    raise ValueError(f"Unsupported rendered report format: {output_path.suffix}")


def _write_docx_report(report: RiskReport, output_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(
            "DOCX export requires python-docx. Install it in the Python environment used for export."
        ) from exc

    top_risks = [_ensure_detail(detail) for detail in report.top3_risks]
    risk_details = [_ensure_detail(detail) for detail in report.risk_details]
    human_review_items = [_ensure_detail(detail) for detail in report.human_review_items]

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal_style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{report.company} 风险分析报告")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"生成时间：{report.timestamp}")
    subtitle_run.font.name = "Arial"
    subtitle_run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    subtitle_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    _add_heading(document, "一、执行摘要", level=1)
    document.add_paragraph(_build_overall_narrative(report))

    summary_table = document.add_table(rows=5, cols=2)
    summary_table.style = "Table Grid"
    summary_pairs = [
        ("综合风险等级", LEVEL_LABELS.get(report.overall_risk_level, report.overall_risk_level)),
        ("综合风险分数", str(report.overall_score)),
        ("模型综合置信度", str(report.confidence)),
        ("风险趋势", TREND_LABELS.get(report.trend, report.trend or "暂无")),
        ("重点风险类别", "、".join(report.top_categories) if report.top_categories else "暂无"),
    ]
    for row, (label, value) in zip(summary_table.rows, summary_pairs):
        row.cells[0].text = label
        row.cells[1].text = value

    _add_heading(document, "二、置信度与红线校验", level=1)
    for line in _build_confidence_lines(report):
        document.add_paragraph(line.removeprefix("- "))

    _add_heading(document, "三、重点风险概览", level=1)
    if top_risks:
        for index, detail in enumerate(top_risks, start=1):
            _add_heading(document, f"Top {index}：{detail.category} / {detail.subtype}", level=2)
            document.add_paragraph(
                f"严重度：{LEVEL_LABELS.get(detail.severity, detail.severity)}；"
                f"风险分数：{detail.score}；"
                f"置信度：{detail.confidence}"
            )
            document.add_paragraph(f"核心证据：{detail.evidence}")
            document.add_paragraph(f"建议措施：{detail.suggestion or '建议结合业务流程进一步制定整改方案。'}")
    else:
        document.add_paragraph("当前未识别出明确风险。")

    _add_heading(document, "四、按生命周期阶段的风险分析与解决方案", level=1)
    if report.lifecycle_stage_groups:
        running_index = 1
        for group in report.lifecycle_stage_groups:
            _add_heading(document, group.stage, level=2)
            document.add_paragraph(group.summary)
            for hint in group.propagation_hints:
                document.add_paragraph(hint, style="List Bullet")
            for detail in group.risk_details:
                _add_heading(document, f"{running_index}. {detail.category} / {detail.subtype}", level=2)
                document.add_paragraph(f"风险分析：{_detail_analysis_text(detail)}")
                document.add_paragraph(f"严重度：{LEVEL_LABELS.get(detail.severity, detail.severity)}")
                document.add_paragraph(f"风险分数：{detail.score}")
                document.add_paragraph(f"置信度：{detail.confidence}")
                document.add_paragraph(f"所属阶段：{detail.lifecycle_stage_hint or '场景落地'}")
                document.add_paragraph(f"原文证据：{detail.evidence}")
                document.add_paragraph(f"解决方案：{detail.suggestion or '建议后续根据具体业务场景补充整改措施。'}")
                document.add_paragraph(f"法规依据：{', '.join(detail.legal_basis) if detail.legal_basis else '暂无'}")
                if detail.revision_reason:
                    document.add_paragraph(f"修正说明：{detail.revision_reason}")
                running_index += 1
    elif risk_details:
        for index, detail in enumerate(risk_details, start=1):
            _add_heading(document, f"{index}. {detail.category} / {detail.subtype}", level=2)
            document.add_paragraph(f"风险分析：{_detail_analysis_text(detail)}")
            document.add_paragraph(f"严重度：{LEVEL_LABELS.get(detail.severity, detail.severity)}")
            document.add_paragraph(f"风险分数：{detail.score}")
            document.add_paragraph(f"置信度：{detail.confidence}")
            document.add_paragraph(f"所属阶段：{detail.lifecycle_stage_hint or '场景落地'}")
            document.add_paragraph(f"原文证据：{detail.evidence}")
            document.add_paragraph(f"解决方案：{detail.suggestion or '建议后续根据具体业务场景补充整改措施。'}")
            document.add_paragraph(f"法规依据：{', '.join(detail.legal_basis) if detail.legal_basis else '暂无'}")
            if detail.revision_reason:
                document.add_paragraph(f"修正说明：{detail.revision_reason}")
    else:
        document.add_paragraph("当前无风险明细。")

    _add_heading(document, "五、建议行动计划", level=1)
    for line in _build_action_plan_lines(report):
        document.add_paragraph(line.removeprefix("- "), style="List Bullet")

    _add_heading(document, "六、人工复核建议", level=1)
    if human_review_items:
        for item in human_review_items:
            document.add_paragraph(
                f"{item.category} / {item.subtype}：建议人工复核，当前置信度为 {item.confidence}。",
                style="List Bullet",
            )
    else:
        document.add_paragraph("当前无必须人工复核的条目。")

    document.save(str(output_path))


def _add_heading(document, text: str, level: int) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    heading = document.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(31, 78, 121)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(17, 24, 39)


def _build_overall_narrative(report: RiskReport) -> str:
    if not report.risk_details:
        return "本次评估未识别出明确的高相关风险条目，建议结合更多业务材料继续核查。"
    highest = _ensure_detail(report.top3_risks[0] if report.top3_risks else report.risk_details[0])
    return (
        f"系统判定该企业当前整体风险水平为"
        f"{LEVEL_LABELS.get(report.overall_risk_level, report.overall_risk_level)}，"
        f"最值得优先关注的是“{highest.category} / {highest.subtype}”。"
        f"从现有证据看，主要问题集中在{_top_categories_text(report)}，"
        "建议优先处理高分风险，再逐步完善中低风险的制度与合规措施。"
    )


def _build_confidence_lines(report: RiskReport) -> list[str]:
    if not report.confidence_breakdown:
        return ["- 当前未提供独立的置信度分解结果。"]

    breakdown = report.confidence_breakdown
    lines = [
        f"- 连续置信度分数：`{breakdown.confidence_score}`",
        f"- 信号强度：`{breakdown.signal_strength}`；鲁棒性：`{breakdown.robustness}`；跨Agent一致性：`{breakdown.cross_agent_consistency}`",
    ]
    if breakdown.disagreement_flags:
        lines.append(f"- 分歧标记：{'；'.join(breakdown.disagreement_flags)}")
    if report.gate_flags.triggered_reasons:
        lines.append(f"- 红线触发：{'；'.join(report.gate_flags.triggered_reasons)}")
    else:
        lines.append("- 红线触发：当前未触发隐私/合法性或伦理/公平性二元红线。")
    return lines


def _top_categories_text(report: RiskReport) -> str:
    if not report.top_categories:
        return "暂无明确分类"
    if len(report.top_categories) == 1:
        return report.top_categories[0]
    return "、".join(report.top_categories)


def _render_top_risk_markdown(index: int, detail: RiskReportDetail) -> list[str]:
    return [
        f"### Top {index}：{detail.category} / {detail.subtype}",
        "",
        f"- 严重度：`{LEVEL_LABELS.get(detail.severity, detail.severity)}`",
        f"- 风险分数：`{detail.score}`",
        f"- 置信度：`{detail.confidence}`",
        f"- 核心证据：{detail.evidence}",
        f"- 建议措施：{detail.suggestion or '建议结合业务流程进一步制定整改方案。'}",
        "",
    ]


def _render_detail_markdown(index: int, detail: RiskReportDetail) -> list[str]:
    return [
        f"### {index}. {detail.category} / {detail.subtype}",
        "",
        f"- 严重度：`{LEVEL_LABELS.get(detail.severity, detail.severity)}`",
        f"- 风险分数：`{detail.score}`",
        f"- 置信度：`{detail.confidence}`",
        f"- 风险分析：{_detail_analysis_text(detail)}",
        f"- 原文证据：{detail.evidence}",
        f"- 解决方案：{detail.suggestion or '建议后续根据具体业务场景补充整改措施。'}",
        f"- 法规依据：{', '.join(detail.legal_basis) if detail.legal_basis else '暂无'}",
        f"- 备注：{detail.revision_reason or '无'}",
        "",
    ]


def _render_detail_html(index: int, detail: RiskReportDetail) -> list[str]:
    return [
        f'<div class="risk-card {escape(detail.severity)}">',
        f"<h3>{index}. {escape(detail.category)} / {escape(detail.subtype)}</h3>",
        f"<p><strong>严重度：</strong>{escape(LEVEL_LABELS.get(detail.severity, detail.severity))}</p>",
        f"<p><strong>风险分数：</strong>{detail.score}</p>",
        f"<p><strong>置信度：</strong>{detail.confidence}</p>",
        f"<p><strong>风险分析：</strong>{escape(_detail_analysis_text(detail))}</p>",
        f"<p><strong>原文证据：</strong>{escape(detail.evidence)}</p>",
        f"<p><strong>解决方案：</strong>{escape(detail.suggestion or '建议后续根据具体业务场景补充整改措施。')}</p>",
        f"<p><strong>法规依据：</strong>{escape(', '.join(detail.legal_basis) if detail.legal_basis else '暂无')}</p>",
        f"<p><strong>备注：</strong>{escape(detail.revision_reason or '无')}</p>",
        "</div>",
    ]


def _detail_analysis_text(detail: RiskReportDetail) -> str:
    base = (
        f"系统认为该条证据反映出“{detail.category}”下的“{detail.subtype}”问题，"
        f"当前严重度为{LEVEL_LABELS.get(detail.severity, detail.severity)}。"
        "如果不及时处理，可能进一步影响企业的合规、运营或声誉表现。"
    )
    if detail.cross_agent_disagreement:
        return base + "该条目同时存在跨 agent 分歧，建议优先人工复核。"
    return base


def _build_action_plan_lines(report: RiskReport) -> list[str]:
    risk_details = [_ensure_detail(detail) for detail in report.risk_details]
    top_risks = [_ensure_detail(detail) for detail in report.top3_risks]
    if not risk_details:
        return ["建议继续补充企业材料，并建立定期复评机制。"]

    high_items = [item for item in risk_details if item.severity == "high"]
    medium_items = [item for item in risk_details if item.severity == "medium"]
    low_items = [item for item in risk_details if item.severity == "low"]
    lines: list[str] = []

    if high_items:
        lines.append("- 立即整改：优先处理高严重度风险，建立专项负责人和整改时限。")
    if medium_items:
        lines.append("- 短期优化：针对中等风险补齐制度、审批流程和证据留痕机制。")
    if low_items:
        lines.append("- 持续治理：对低风险项建立常态化复核与制度完善计划。")

    suggestions = []
    for item in top_risks:
        if item.suggestion and item.suggestion not in suggestions:
            suggestions.append(item.suggestion)
    lines.extend([f"- 重点建议：{suggestion}" for suggestion in suggestions])
    return lines or ["建议继续补充企业材料，并建立定期复评机制。"]


def _ensure_detail(detail: RiskReportDetail | dict) -> RiskReportDetail:
    if isinstance(detail, RiskReportDetail):
        return detail
    return RiskReportDetail(**detail)
